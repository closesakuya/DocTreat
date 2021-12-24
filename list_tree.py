import re
from pathlib import Path
from pathlib import WindowsPath
from typing import Optional, List
import os



class DirectionTree:
    def __init__(self,
                 direction_name: str = 'WorkingDirection',
                 direction_path: str = '.',
                 ignore_list: Optional[List[str]] = None,
                 show_cnt: bool = False,
                 use_to_word: bool = False):
        self.owner: WindowsPath = Path(direction_path)
        self.tree: str = direction_name + '/\n'
        self.ignore_list = ignore_list
        self.show_cnt = show_cnt
        self.use_to_word = use_to_word
        if ignore_list is None:
            self.ignore_list = []
        self.direction_ergodic(path_object=self.owner, n=0)

    def tree_add(self, path_object: WindowsPath, n=0, last=False, cnt=0, enum_str=""):
        cnt_str = ""
        if self.show_cnt:
            cnt_str = enum_str + str(cnt+1) + ". "
            cnt_str = cnt_str.replace("..", ".").strip(".")
        if n > 0:
            if last:
                if self.use_to_word:
                    self.tree += '│' + ('        │' * (n - 1)) + '        └────' + cnt_str + path_object.name
                else:
                    self.tree += '│' + ('    │' * (n - 1)) + '    └────' + cnt_str + path_object.name
            else:
                if self.use_to_word:
                    self.tree += '│' + ('        │' * (n - 1)) + '        ├────' + cnt_str + path_object.name
                else:
                    self.tree += '│' + ('    │' * (n - 1)) + '    ├────' + cnt_str + path_object.name
        else:
            if last:
                self.tree += '└' + ('──' * 2) + cnt_str + path_object.name
            else:
                self.tree += '├' + ('──' * 2) + cnt_str + path_object.name
        if path_object.is_file():
            self.tree += '\n'
            return False
        elif path_object.is_dir():
            self.tree += '/\n'
            return True

    def filter_file(self, file):
        for item in self.ignore_list:
            if re.fullmatch(item, file.name):
                return False
        return True

    def direction_ergodic(self, path_object: WindowsPath, n=0, enum_str=""):
        dir_file: list = list(path_object.iterdir())
        dir_file.sort(key=lambda x: x.name.lower())
        dir_file = [f for f in filter(self.filter_file, dir_file)]
        for i, item in enumerate(dir_file):
            if i + 1 == len(dir_file):
                if self.tree_add(item, n, last=True, cnt=i, enum_str=enum_str):
                    self.direction_ergodic(item, n + 1, enum_str=enum_str+"." + str(i+1) + ".")
            else:
                if self.tree_add(item, n, last=False, cnt=i, enum_str=enum_str):
                    self.direction_ergodic(item, n + 1, enum_str=enum_str+"." + str(i+1) + ".")


def list_tree(src: str, dst: str = None, ignore_lst: List[str] = None, logger = None, use_to_word: bool = True):
    tree = DirectionTree(ignore_list=ignore_lst, direction_path=src,
                         direction_name=os.path.splitext(os.path.split(src)[-1])[0],
                         show_cnt=True, use_to_word=use_to_word)
    # print(dst)
    lst = tree.tree.__str__()
    if callable(logger):
        logger(lst)
    if dst:
        suffix = os.path.splitext(os.path.split(dst)[-1])[-1]
        if suffix == ".doc" or suffix == ".docx":
            import docx
            from docx.shared import Pt
            document = docx.Document()
            from docx.oxml.ns import qn

            mystyle = document.styles["Normal"]
            font = mystyle.font
            # 初号=42磅 小初=36磅 一号=26磅 小一=24磅?二号=22磅 小二=18磅 三号=16磅
            # 小三=15磅?四号=14磅 小四=12磅 五号=10.5磅 小五=9磅?六号=7.5磅 小六=6.5磅 七号=5.5磅 八号=5磅
            font.size = Pt(9)  # 小五
            font.name = u'宋体'
            font._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            # 添加标题，并设置级别，范围：0 至 9，默认为1
            p = document.add_heading('{0}目录结构'.format(
                os.path.splitext(os.path.split(src)[-1])[0]), 0)

            # 添加段落，文本可以包含制表符（\t）、换行符（\n）或回车符（\r）等
            for item in lst.split("\n"):
                blank = ''
                for s in item:
                    if s == '│' or s == '├' or s == '└':
                        blank += '│'
                    else:
                        blank += ' '
                p = document.add_paragraph(blank)
                p.style = mystyle
                paragraph_format = p.paragraph_format
                paragraph_format.space_before = Pt(0)  # 上行间距
                paragraph_format.space_after = Pt(0)  # 下行间距
                paragraph_format.line_spacing = Pt(12)  # 行距
                p = document.add_paragraph(item)
                p.style = mystyle
                paragraph_format = p.paragraph_format
                paragraph_format.space_before = Pt(0)  # 上行间距
                paragraph_format.space_after = Pt(0)  # 下行间距
                paragraph_format.line_spacing = Pt(12)  # 行距

            document.save(dst)





# if __name__ == '__main__':
#     i_l = [
#         '\.git', '__pycache__', 'test.+', 'venv', '.+\.whl', '\.idea', '.+\.jpg', '.+\.png',
#         'image', 'css', 'admin', 'tool.py', 'db.sqlite3',
#         '~\$.*\.xls[x]?', '~\$.*\.doc[x]?', '~\$.*\.ppt[x]?', '~.*\..*'
#     ]
#
#     list_tree(r"C:\Users\LXW\Desktop\界面调研", ignore_lst=i_l,
#               dst=os.getcwd() + os.sep + "result.docx", logger=lambda x: print(x))